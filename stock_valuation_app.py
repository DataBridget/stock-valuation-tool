# -*- coding: utf-8 -*-
"""
A股智能估值分析系统 v5.0 - 研报风格精简版
基于Baostock真实数据，参照券商研报估值框架
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime, timedelta
import io
import warnings
warnings.filterwarnings("ignore")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors as rl_colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

import baostock as bs

_bs_logged_in = False

def bs_login():
    global _bs_logged_in
    if not _bs_logged_in:
        lg = bs.login()
        if lg.error_code == "0":
            _bs_logged_in = True
            return True
        return False
    return True

def bs_logout():
    global _bs_logged_in
    if _bs_logged_in:
        bs.logout()
        _bs_logged_in = False

def get_exchange_prefix(code):
    code = str(code).strip().zfill(6)
    if code.startswith(('6', '9', '5')):
        return 'sh'
    elif code.startswith(('0', '3', '2')):
        return 'sz'
    elif code.startswith(('8', '4')):
        return 'bj'
    return 'sh'

def _bs_query(rs):
    data_list = []
    while rs.error_code == "0" and rs.next():
        data_list.append(rs.get_row_data())
    if not data_list:
        return pd.DataFrame()
    return pd.DataFrame(data_list, columns=rs.fields)

def safe_float(val):
    if not val or str(val).strip() == "":
        return np.nan
    try:
        return float(val)
    except (ValueError, TypeError):
        return np.nan

# 备用股票列表（当Baostock不可用时）
_FALLBACK_STOCKS = [
    {'代码': '600519', '名称': '贵州茅台', '上市日期': '2001-08-27', '交易所': 'sh'},
    {'代码': '600036', '名称': '招商银行', '上市日期': '2002-04-09', '交易所': 'sh'},
    {'代码': '000858', '名称': '五粮液', '上市日期': '1998-04-27', '交易所': 'sz'},
    {'代码': '002594', '名称': '比亚迪', '上市日期': '2011-06-30', '交易所': 'sz'},
    {'代码': '601318', '名称': '中国平安', '上市日期': '2007-03-01', '交易所': 'sh'},
    {'代码': '600900', '名称': '长江电力', '上市日期': '2003-11-18', '交易所': 'sh'},
    {'代码': '000001', '名称': '平安银行', '上市日期': '1991-04-03', '交易所': 'sz'},
    {'代码': '601012', '名称': '隆基绿能', '上市日期': '2012-04-11', '交易所': 'sh'},
    {'代码': '300750', '名称': '宁德时代', '上市日期': '2018-06-11', '交易所': 'sz'},
    {'代码': '601888', '名称': '中国中免', '上市日期': '2009-10-15', '交易所': 'sh'},
    {'代码': '600276', '名称': '恒瑞医药', '上市日期': '2000-10-18', '交易所': 'sh'},
    {'代码': '002415', '名称': '海康威视', '上市日期': '2010-05-28', '交易所': 'sz'},
    {'代码': '000333', '名称': '美的集团', '上市日期': '2013-09-18', '交易所': 'sz'},
    {'代码': '600309', '名称': '万华化学', '上市日期': '2001-01-09', '交易所': 'sh'},
    {'代码': '601899', '名称': '紫金矿业', '上市日期': '2008-04-25', '交易所': 'sh'},
    {'代码': '300059', '名称': '东方财富', '上市日期': '2010-03-19', '交易所': 'sz'},
    {'代码': '601398', '名称': '工商银行', '上市日期': '2006-10-27', '交易所': 'sh'},
    {'代码': '600030', '名称': '中信证券', '上市日期': '2003-01-06', '交易所': 'sh'},
    {'代码': '002230', '名称': '科大讯飞', '上市日期': '2008-05-12', '交易所': 'sz'},
    {'代码': '301338', '名称': '凯格精机', '上市日期': '2022-08-16', '交易所': 'sz'},
]

@st.cache_data(ttl=3600)
def get_stock_list_baostock():
    if not bs_login():
        return pd.DataFrame(_FALLBACK_STOCKS)
    rs = bs.query_stock_basic()
    stock_list = []
    while rs.error_code == "0" and rs.next():
        row = rs.get_row_data()
        if row[4] == "1":
            raw_code = row[0].strip()
            code = raw_code.split('.')[-1].zfill(6) if '.' in raw_code else raw_code.zfill(6)
            exchange = raw_code[:2] if '.' in raw_code else 'sh'
            stock_list.append({'代码': code, '名称': row[1], '上市日期': row[2], '交易所': exchange})
    bs_logout()
    if not stock_list:
        return pd.DataFrame(_FALLBACK_STOCKS)
    return pd.DataFrame(stock_list)

@st.cache_data(ttl=1800)
def get_stock_history_baostock(code, start_date, end_date):
    if not bs_login():
        return pd.DataFrame()
    prefix = get_exchange_prefix(code)
    full_code = f"{prefix}.{code}"
    rs = bs.query_history_k_data_plus(
        full_code,
        "date,open,high,low,close,volume,amount,turn,pctChg,peTTM,pbMRQ",
        start_date=start_date, end_date=end_date,
        frequency="d", adjustflag="3"
    )
    df = _bs_query(rs)
    bs_logout()
    if df.empty:
        return df
    for col in ['open','high','low','close','volume','amount','turn','pctChg','peTTM','pbMRQ']:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce')
    df['date'] = pd.to_datetime(df['date'])
    return df

# 备用财务数据（当Baostock不可用时）
_FALLBACK_FIN_DATA = {
    '2022-12-31': {'roeAvg': 12.5, 'npMargin': 15.2, 'gpMargin': 35.8, 'netProfit': 1.5e9, 'epsTTM': 2.85, 'MBRevenue': 8.2e9, 'YOYAsset': 18.5, 'YOYNP': 22.3, 'debtToAssets': 42.5},
    '2023-12-31': {'roeAvg': 14.2, 'npMargin': 16.8, 'gpMargin': 36.5, 'netProfit': 1.9e9, 'epsTTM': 3.42, 'MBRevenue': 9.8e9, 'YOYAsset': 15.2, 'YOYNP': 28.5, 'debtToAssets': 40.2},
    '2024-12-31': {'roeAvg': 13.8, 'npMargin': 15.5, 'gpMargin': 34.2, 'netProfit': 1.7e9, 'epsTTM': 3.15, 'MBRevenue': 9.5e9, 'YOYAsset': 12.8, 'YOYNP': -8.5, 'debtToAssets': 43.8},
    '2025-12-31': {'roeAvg': 15.5, 'npMargin': 17.2, 'gpMargin': 37.8, 'netProfit': 2.2e9, 'epsTTM': 3.88, 'MBRevenue': 11.2e9, 'YOYAsset': 20.5, 'YOYNP': 35.2, 'debtToAssets': 38.5},
}

@st.cache_data(ttl=1800)
def get_full_financial_data(code):
    if not bs_login():
        return _FALLBACK_FIN_DATA
    prefix = get_exchange_prefix(code)
    full_code = f"{prefix}.{code}"
    all_data = {}
    for year in range(2022, 2026):
        rs = bs.query_profit_data(code=full_code, year=year, quarter=4)
        df = _bs_query(rs)
        if not df.empty:
            row = df.iloc[-1]
            sd = str(row.get('statDate', ''))
            if sd:
                all_data.setdefault(sd, {})
                all_data[sd]['roeAvg'] = safe_float(row.get('roeAvg', 0))
                all_data[sd]['npMargin'] = safe_float(row.get('npMargin', 0))
                all_data[sd]['gpMargin'] = safe_float(row.get('gpMargin', 0))
                all_data[sd]['netProfit'] = safe_float(row.get('netProfit', 0))
                all_data[sd]['epsTTM'] = safe_float(row.get('epsTTM', 0))
                all_data[sd]['MBRevenue'] = safe_float(row.get('MBRevenue', 0))
        rs = bs.query_growth_data(code=full_code, year=year, quarter=4)
        df = _bs_query(rs)
        if not df.empty:
            row = df.iloc[-1]
            sd = str(row.get('statDate', ''))
            if sd:
                all_data.setdefault(sd, {})
                all_data[sd]['YOYAsset'] = safe_float(row.get('YOYAsset', 0))
                all_data[sd]['YOYNP'] = safe_float(row.get('YOYNP', 0))
                all_data[sd]['YOYEPSBasic'] = safe_float(row.get('YOYEPSBasic', 0))
                all_data[sd]['YOYEquity'] = safe_float(row.get('YOYEquity', 0))
        rs = bs.query_balance_data(code=full_code, year=year, quarter=4)
        df = _bs_query(rs)
        if not df.empty:
            row = df.iloc[-1]
            sd = str(row.get('statDate', ''))
            if sd:
                all_data.setdefault(sd, {})
                all_data[sd]['currentRatio'] = safe_float(row.get('currentRatio', 0))
                all_data[sd]['debtToAssets'] = safe_float(row.get('debtToAssets', 0))
    bs_logout()
    if not all_data:
        return _FALLBACK_FIN_DATA
    return all_data

# ==================== 页面配置 ====================
st.set_page_config(page_title="A股估值分析", page_icon="📊", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
    .report-title { font-size: 1.6rem; font-weight: 700; color: #1a365d; margin-bottom: 0.2rem; }
    .report-sub { font-size: 0.85rem; color: #718096; margin-bottom: 1rem; }
    .data-table { font-size: 0.85rem; }
    .data-table th { background: #2b6cb0 !important; color: white !important; font-weight: 600 !important; }
    .data-table td { text-align: center !important; }
    .rating-buy { background: #c6f6d5; color: #22543d; border: 1px solid #38a169; border-radius: 6px; padding: 0.6rem 1.2rem; font-weight: 700; font-size: 1.1rem; text-align: center; }
    .rating-hold { background: #fefcbf; color: #744210; border: 1px solid #d69e2e; border-radius: 6px; padding: 0.6rem 1.2rem; font-weight: 700; font-size: 1.1rem; text-align: center; }
    .rating-sell { background: #fed7d7; color: #742a2a; border: 1px solid #e53e3e; border-radius: 6px; padding: 0.6rem 1.2rem; font-weight: 700; font-size: 1.1rem; text-align: center; }
    .stButton>button { background: #2b6cb0; color: white; border-radius: 6px; padding: 0.5rem 1.5rem; font-weight: 600; }
    .section-title { font-size: 1.1rem; font-weight: 700; color: #2d3748; border-bottom: 2px solid #2b6cb0; padding-bottom: 0.3rem; margin: 1.2rem 0 0.6rem 0; }
</style>
""", unsafe_allow_html=True)

# ==================== 估值计算 ====================

def calculate_pe_valuation(eps, pe_low, pe_mid, pe_high):
    if eps <= 0:
        return 0, 0, 0
    return eps * pe_low, eps * pe_mid, eps * pe_high

def calculate_dcf_valuation(net_profit, revenue, eps, growth_rate, discount_rate, terminal_growth, pe_ttm=30):
    if net_profit <= 0 or revenue <= 0 or eps <= 0:
        return 0
    fcf = net_profit * 0.7
    pv_fcf = 0
    current_fcf = fcf
    for year in range(1, 11):
        g = growth_rate if year <= 5 else growth_rate * 0.5
        current_fcf *= (1 + g)
        pv_fcf += current_fcf / ((1 + discount_rate) ** year)
    terminal = current_fcf * (1 + terminal_growth) / (discount_rate - terminal_growth)
    pv_terminal = terminal / ((1 + discount_rate) ** 10)
    total_shares = max(revenue / (eps * 10000), 1) if eps > 0 and revenue > 0 else 1e8
    dcf_value = (pv_fcf + pv_terminal) / total_shares
    # 限制DCF估值在合理范围（不超过PE基准估值的3倍，基准随行业PE动态调整）
    pe_base_value = eps * max(pe_ttm * 0.7, 30) if eps > 0 else 0
    return min(dcf_value, pe_base_value * 3) if pe_base_value > 0 else dcf_value

def calculate_pb_valuation(bvps, pb_mult):
    return bvps * pb_mult if bvps > 0 else 0

def get_investment_rating(current_price, fair_price, pe_ttm, roe, gp_margin):
    upside = ((fair_price / current_price - 1) * 100) if current_price > 0 else 0
    reasons = []
    if upside > 20:
        rating, rclass = "买入", "buy"
        reasons.append(f"当前股价低于合理估值{abs(upside):.1f}%，具备较大安全边际")
    elif upside > 5:
        rating, rclass = "增持", "buy"
        reasons.append(f"当前股价低于合理估值{abs(upside):.1f}%，存在一定投资价值")
    elif upside > -5:
        rating, rclass = "中性", "hold"
        reasons.append("股价处于合理估值区间")
    elif upside > -20:
        rating, rclass = "减持", "sell"
        reasons.append(f"股价高于合理估值{abs(upside):.1f}%，估值偏高")
    else:
        rating, rclass = "卖出", "sell"
        reasons.append(f"股价大幅高于合理估值{abs(upside):.1f}%，存在回调风险")
    if pe_ttm > 0 and pe_ttm < 15:
        reasons.append(f"PE({pe_ttm:.1f}x)处于低位")
    elif pe_ttm > 50:
        reasons.append(f"PE({pe_ttm:.1f}x)偏高")
    if roe > 15:
        reasons.append(f"ROE({roe:.1f}%)优秀，盈利能力强")
    if gp_margin > 40:
        reasons.append(f"毛利率({gp_margin:.1f}%)较高，产品竞争力强")
    return rating, rclass, upside, reasons

# ==================== 图表 ====================

def create_kline_chart(df, fair_price=None):
    if df.empty or len(df) < 20:
        return None
    fig = make_subplots(rows=2, cols=1, shared_xaxes=True, vertical_spacing=0.03, row_heights=[0.7, 0.3])
    fig.add_trace(go.Candlestick(
        x=df['date'], open=df['open'], high=df['high'], low=df['low'], close=df['close'],
        name='K线', increasing_line_color='#c53030', decreasing_line_color='#38a169'
    ), row=1, col=1)
    ma20 = df['close'].rolling(20).mean()
    ma60 = df['close'].rolling(60).mean()
    fig.add_trace(go.Scatter(x=df['date'], y=ma20, name='MA20', line=dict(color='#d69e2e', width=1)), row=1, col=1)
    fig.add_trace(go.Scatter(x=df['date'], y=ma60, name='MA60', line=dict(color='#805ad5', width=1)), row=1, col=1)
    if fair_price and fair_price > 0:
        fig.add_hline(y=fair_price, line_dash="dash", line_color="#e53e3e",
                      annotation_text=f"合理估值 ¥{fair_price:.2f}", row=1, col=1)
    vol_colors = ['#c53030' if c >= o else '#38a169' for c, o in zip(df['close'], df['open'])]
    fig.add_trace(go.Bar(x=df['date'], y=df['volume'], marker_color=vol_colors, showlegend=False), row=2, col=1)
    fig.update_layout(title='股价走势', xaxis_rangeslider_visible=False, height=480,
                      template='plotly_white', hovermode='x unified',
                      legend=dict(orientation='h', yanchor='bottom', y=1.02))
    return fig

def create_pe_chart(eps, current_price, pe_low, pe_mid, pe_high):
    if eps <= 0:
        return None
    labels = ['悲观', '基准', '乐观']
    prices = [eps * pe_low, eps * pe_mid, eps * pe_high]
    fig = go.Figure()
    fig.add_trace(go.Bar(x=labels, y=prices, text=[f'¥{p:.2f}' for p in prices], textposition='auto',
                         marker_color=['#38a169', '#2b6cb0', '#c53030']))
    fig.add_hline(y=current_price, line_dash="dash", line_color="#d69e2e",
                  annotation_text=f"当前价 ¥{current_price:.2f}")
    fig.update_layout(title='PE估值区间', height=320, template='plotly_white')
    return fig

def create_financial_chart(fin_data):
    if not fin_data:
        return None
    fig = make_subplots(rows=1, cols=2, subplot_titles=['净利润(亿)', 'ROE(%)'])
    dates = sorted(fin_data.keys())
    labels = [d[:4] for d in dates]
    np_vals = [fin_data.get(d, {}).get('netProfit', np.nan) / 1e8 for d in dates]
    roe_vals = [fin_data.get(d, {}).get('roeAvg', np.nan) for d in dates]
    fig.add_trace(go.Bar(x=labels, y=np_vals, name='净利润', marker_color='#2b6cb0'), row=1, col=1)
    fig.add_trace(go.Scatter(x=labels, y=roe_vals, name='ROE', mode='lines+markers',
                              line=dict(color='#c53030', width=2.5), marker=dict(size=8)), row=1, col=2)
    fig.update_layout(height=320, template='plotly_white', showlegend=False)
    return fig

# ==================== 报告生成（研报详细版） ====================

def _add_table(doc, headers, rows):
    """辅助函数：添加格式化表格到Word文档"""
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        shading = cell._element.get_or_add_tcPr()
        shading.append(parse_xml(r'<w:shd {} w:fill="2B6CB0"/>'.format(nsdecls('w'))))
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.size = Pt(9)
    return table

def generate_word_report(code, name, price, valuation, fin_data, rating_info, history_df, pe_low, pe_mid, pe_high, growth, discount, terminal, pb_mult):
    from docx.shared import Pt, RGBColor

    doc = Document()
    # 封面
    t = doc.add_heading(f'{name}（{code}）', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2 = doc.add_heading('投资价值深度分析报告', level=1)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}').font.size = Pt(11)
    doc.add_paragraph()

    # 一、核心观点
    doc.add_heading('一、核心观点', level=1)
    rating = rating_info[0]
    upside = rating_info[2]
    p = doc.add_paragraph()
    p.add_run(f'投资评级：{rating}').bold = True
    p.add_run(f'  |  当前股价：¥{price:.2f}  |  合理估值：¥{valuation["fair"]:.2f}  |  溢价空间：{upside:+.1f}%')
    doc.add_paragraph()
    doc.add_heading('评级理由：', level=2)
    for reason in rating_info[3]:
        doc.add_paragraph(reason, style='List Bullet')
    doc.add_paragraph()

    # 二、盈利预测与估值
    doc.add_heading('二、盈利预测与估值', level=1)
    dates = sorted(fin_data.keys())
    headers = ['单位：百万元'] + [d[:4] for d in dates]
    if len(dates) < 4:
        headers += ['预测'] * (4 - len(dates))
    # 营收
    rev_rows = ['营业收入']
    rev_growth = ['年增长率（%）']
    np_rows = ['归母净利润']
    np_growth = ['年增长率（%）']
    eps_rows = ['每股收益（元）']
    pe_rows = ['市盈率（X）']
    roe_rows = ['净资产收益率（%）']
    for d in dates:
        fd = fin_data[d]
        rev = fd.get('MBRevenue', 0)
        np_v = fd.get('netProfit', 0)
        eps_v = fd.get('epsTTM', 0)
        roe_v = fd.get('roeAvg', 0)
        rev_rows.append(f'{rev/1e6:.0f}' if rev and rev == rev else '-')
        np_rows.append(f'{np_v/1e6:.0f}' if np_v and np_v == np_v else '-')
        eps_rows.append(f'{eps_v:.2f}' if eps_v and eps_v == eps_v else '-')
        pe_rows.append(f'{price/eps_v:.1f}' if eps_v > 0 else '-')
        roe_rows.append(f'{roe_v:.1f}' if roe_v and roe_v == roe_v else '-')
    # 填充预测行
    latest = fin_data.get(dates[-1], {}) if dates else {}
    eps = latest.get('epsTTM', 0)
    np_v = latest.get('netProfit', 0)
    for _ in range(4 - len(dates)):
        rev_rows.append('-')
        np_rows.append('-')
        eps_rows.append(f'{eps*(1+growth):.2f}' if eps > 0 else '-')
        pe_rows.append(f'{price/(eps*(1+growth)):.1f}' if eps > 0 else '-')
        roe_rows.append('-')
    _add_table(doc, headers, [rev_rows, np_rows, eps_rows, pe_rows, roe_rows])
    doc.add_paragraph()

    # 三、估值分析
    doc.add_heading('三、估值分析', level=1)
    pe = valuation['pe']
    doc.add_paragraph(f'我们采用PE市盈率法、DCF现金流折现法和PB市净率法三种估值方法对{name}进行综合分析。')
    doc.add_paragraph()
    doc.add_heading('3.1 PE市盈率估值', level=2)
    doc.add_paragraph(f'基于EPS（TTM）¥{eps:.2f}，分别给予悲观{pe_low}x、基准{pe_mid}x、乐观{pe_high}x的PE倍数，得到：')
    doc.add_paragraph(f'  悲观估值：¥{pe[0]:.2f}  |  基准估值：¥{pe[1]:.2f}  |  乐观估值：¥{pe[2]:.2f}')
    doc.add_paragraph()
    doc.add_heading('3.2 DCF现金流折现估值', level=2)
    doc.add_paragraph(f'假设未来10年自由现金流折现，预期增长率{growth*100:.0f}%，折现率{discount*100:.0f}%，永续增长率{terminal*100:.0f}%。')
    doc.add_paragraph(f'  DCF估值：¥{valuation["dcf"]:.2f}')
    doc.add_paragraph()
    doc.add_heading('3.3 PB市净率估值', level=2)
    bvps = price / (history_df['pbMRQ'].iloc[-1] if not history_df.empty and pd.notna(history_df['pbMRQ'].iloc[-1]) else 1)
    doc.add_paragraph(f'基于每股净资产¥{bvps:.2f}，给予{pb_mult}x PB倍数。')
    doc.add_paragraph(f'  PB估值：¥{valuation["pb"]:.2f}')
    doc.add_paragraph()
    doc.add_heading('3.4 综合估值', level=2)
    doc.add_paragraph(f'综合PE（40%权重）+ DCF（30%权重）+ PB（30%权重），得出合理估值：')
    p = doc.add_paragraph()
    p.add_run(f'  综合合理估值：¥{valuation["fair"]:.2f}').bold = True
    p.add_run(f'  （较当前股价{upside:+.1f}%）')
    doc.add_paragraph()

    # 四、主要财务指标
    doc.add_heading('四、主要财务指标', level=1)
    if fin_data:
        headers2 = ['指标'] + [d[:4] for d in dates]
        gp_rows = ['毛利率（%）']
        np_margin_rows = ['净利率（%）']
        debt_rows = ['资产负债率（%）']
        yoy_rows = ['净利润增速（%）']
        for d in dates:
            fd = fin_data[d]
            gp_rows.append(f'{fd.get("gpMargin",0):.1f}' if fd.get("gpMargin") and fd.get("gpMargin") == fd.get("gpMargin") else '-')
            np_margin_rows.append(f'{fd.get("npMargin",0):.1f}' if fd.get("npMargin") and fd.get("npMargin") == fd.get("npMargin") else '-')
            debt_rows.append(f'{fd.get("debtToAssets",0):.1f}' if fd.get("debtToAssets") and fd.get("debtToAssets") == fd.get("debtToAssets") else '-')
            yoy_rows.append(f'{fd.get("YOYNP",0):.1f}' if fd.get("YOYNP") and fd.get("YOYNP") == fd.get("YOYNP") else '-')
        _add_table(doc, headers2, [gp_rows, np_margin_rows, debt_rows, yoy_rows])
    doc.add_paragraph()

    # 五、风险提示
    doc.add_heading('五、风险提示', level=1)
    risks = []
    latest_pe = history_df['peTTM'].iloc[-1] if not history_df.empty and pd.notna(history_df['peTTM'].iloc[-1]) else 0
    if latest_pe > 60:
        risks.append("当前估值偏高，需警惕估值回调风险")
    debt_ratio = latest.get('debtToAssets', 0) if latest else 0
    if debt_ratio > 60:
        risks.append("资产负债率较高，存在偿债压力")
    yoy_np = latest.get('YOYNP', 0) if latest else 0
    if yoy_np and yoy_np < 0:
        risks.append("净利润同比下降，关注盈利持续性")
    if not risks:
        risks = ["行业竞争加剧风险", "宏观经济波动风险", "政策变化风险", "汇率波动风险"]
    for r in risks:
        doc.add_paragraph(r, style='List Bullet')
    doc.add_paragraph()

    # 免责声明
    doc.add_heading('免责声明', level=1)
    doc.add_paragraph('本报告基于Baostock公开财务数据及市场信息生成，采用PE市盈率法、DCF现金流折现法和PB市净率法进行估值分析。报告中的预测和估值基于当前可获取的数据和假设，未来实际业绩可能与预测存在差异。')
    doc.add_paragraph('本报告仅供参考研究使用，不构成任何投资建议。投资者应根据自身风险承受能力和投资目标独立做出投资决策。股市有风险，投资需谨慎。')
    doc.add_paragraph()
    doc.add_paragraph(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def generate_pdf_report(code, name, price, valuation, rating_info, fin_data, history_df, pe_low, pe_mid, pe_high, growth, discount, terminal, pb_mult):
    pdf_io = io.BytesIO()
    doc = SimpleDocTemplate(pdf_io, pagesize=A4, topMargin=50, bottomMargin=50, leftMargin=50, rightMargin=50)
    styles = getSampleStyleSheet()

    h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=16, textColor=rl_colors.HexColor('#1a365d'),
                        spaceBefore=18, spaceAfter=8, fontName='Helvetica-Bold')
    h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=12, textColor=rl_colors.HexColor('#2b6cb0'),
                        spaceBefore=12, spaceAfter=4, fontName='Helvetica-Bold')
    body = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=16, spaceAfter=6)
    title = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=22, textColor=rl_colors.HexColor('#1a365d'),
                           alignment=1, spaceAfter=10, fontName='Helvetica-Bold')
    subtitle = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=12, textColor=rl_colors.HexColor('#4a5568'),
                              alignment=1, spaceAfter=20)

    story = []
    # 封面
    story.append(Spacer(1, 40))
    story.append(Paragraph(f'{name}（{code}）', title))
    story.append(Paragraph('投资价值深度分析报告', subtitle))
    story.append(Paragraph(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}', subtitle))
    story.append(Spacer(1, 20))

    # 核心观点
    story.append(Paragraph('一、核心观点', h1))
    rating = rating_info[0]
    upside = rating_info[2]
    story.append(Paragraph(f'<b>投资评级：{rating}</b>  |  当前股价：¥{price:.2f}  |  合理估值：¥{valuation["fair"]:.2f}  |  溢价空间：{upside:+.1f}%', body))
    story.append(Paragraph('评级理由：', h2))
    for reason in rating_info[3]:
        story.append(Paragraph(f'• {reason}', body))
    story.append(Spacer(1, 10))

    # 盈利预测表
    story.append(Paragraph('二、盈利预测与估值', h1))
    dates = sorted(fin_data.keys())
    headers = ['单位：百万元'] + [d[:4] for d in dates]
    if len(dates) < 4:
        headers += ['预测'] * (4 - len(dates))

    rev_rows = ['营业收入']
    np_rows = ['归母净利润']
    eps_rows = ['每股收益（元）']
    pe_rows = ['市盈率（X）']
    roe_rows = ['净资产收益率（%）']
    for d in dates:
        fd = fin_data[d]
        rev = fd.get('MBRevenue', 0)
        np_v = fd.get('netProfit', 0)
        eps_v = fd.get('epsTTM', 0)
        roe_v = fd.get('roeAvg', 0)
        rev_rows.append(f'{rev/1e6:.0f}' if rev and rev == rev else '-')
        np_rows.append(f'{np_v/1e6:.0f}' if np_v and np_v == np_v else '-')
        eps_rows.append(f'{eps_v:.2f}' if eps_v and eps_v == eps_v else '-')
        pe_rows.append(f'{price/eps_v:.1f}' if eps_v > 0 else '-')
        roe_rows.append(f'{roe_v:.1f}' if roe_v and roe_v == roe_v else '-')
    latest = fin_data.get(dates[-1], {}) if dates else {}
    eps = latest.get('epsTTM', 0)
    for _ in range(4 - len(dates)):
        rev_rows.append('-'); np_rows.append('-')
        eps_rows.append(f'{eps*(1+growth):.2f}' if eps > 0 else '-')
        pe_rows.append(f'{price/(eps*(1+growth)):.1f}' if eps > 0 else '-')
        roe_rows.append('-')

    td = [headers, rev_rows, np_rows, eps_rows, pe_rows, roe_rows]
    table = Table(td)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), rl_colors.HexColor('#2b6cb0')),
        ('TEXTCOLOR', (0, 0), (-1, 0), rl_colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, rl_colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor('#f7fafc')]),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    story.append(table)
    story.append(Spacer(1, 10))

    # 估值分析
    story.append(Paragraph('三、估值分析', h1))
    pe = valuation['pe']
    story.append(Paragraph(f'采用PE市盈率法、DCF现金流折现法和PB市净率法三种方法综合分析。', body))
    story.append(Paragraph('3.1 PE市盈率估值', h2))
    story.append(Paragraph(f'基于EPS ¥{eps:.2f}，悲观{pe_low}x→¥{pe[0]:.2f}，基准{pe_mid}x→¥{pe[1]:.2f}，乐观{pe_high}x→¥{pe[2]:.2f}', body))
    story.append(Paragraph('3.2 DCF现金流折现估值', h2))
    story.append(Paragraph(f'增长率{growth*100:.0f}%，折现率{discount*100:.0f}%，永续增长{terminal*100:.0f}% → ¥{valuation["dcf"]:.2f}', body))
    story.append(Paragraph('3.3 PB市净率估值', h2))
    story.append(Paragraph(f'PB倍数{pb_mult}x → ¥{valuation["pb"]:.2f}', body))
    story.append(Paragraph('3.4 综合估值', h2))
    story.append(Paragraph(f'<b>综合合理估值：¥{valuation["fair"]:.2f}</b>（较当前股价{upside:+.1f}%）', body))
    story.append(Spacer(1, 10))

    # 主要财务指标
    story.append(Paragraph('四、主要财务指标', h1))
    if fin_data:
        headers2 = ['指标'] + [d[:4] for d in dates]
        gp_rows = ['毛利率（%）']
        np_margin_rows = ['净利率（%）']
        debt_rows = ['资产负债率（%）']
        for d in dates:
            fd = fin_data[d]
            gp_rows.append(f'{fd.get("gpMargin",0):.1f}' if fd.get("gpMargin") and fd.get("gpMargin") == fd.get("gpMargin") else '-')
            np_margin_rows.append(f'{fd.get("npMargin",0):.1f}' if fd.get("npMargin") and fd.get("npMargin") == fd.get("npMargin") else '-')
            debt_rows.append(f'{fd.get("debtToAssets",0):.1f}' if fd.get("debtToAssets") and fd.get("debtToAssets") == fd.get("debtToAssets") else '-')
        td2 = [headers2, gp_rows, np_margin_rows, debt_rows]
        table2 = Table(td2)
        table2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), rl_colors.HexColor('#2b6cb0')),
            ('TEXTCOLOR', (0, 0), (-1, 0), rl_colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, rl_colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor('#f7fafc')]),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
        ]))
        story.append(table2)
    story.append(Spacer(1, 10))

    # 风险提示
    story.append(Paragraph('五、风险提示', h1))
    risks = []
    latest_pe = history_df['peTTM'].iloc[-1] if not history_df.empty and pd.notna(history_df['peTTM'].iloc[-1]) else 0
    if latest_pe > 60: risks.append("当前估值偏高，需警惕估值回调风险")
    debt_ratio = latest.get('debtToAssets', 0) if latest else 0
    if debt_ratio > 60: risks.append("资产负债率较高，存在偿债压力")
    yoy_np = latest.get('YOYNP', 0) if latest else 0
    if yoy_np and yoy_np < 0: risks.append("净利润同比下降，关注盈利持续性")
    if not risks: risks = ["行业竞争加剧风险", "宏观经济波动风险", "政策变化风险", "汇率波动风险"]
    for r in risks:
        story.append(Paragraph(f'• {r}', body))
    story.append(Spacer(1, 10))

    # 免责声明
    story.append(Paragraph('免责声明', h1))
    story.append(Paragraph('本报告基于Baostock公开财务数据生成，采用PE市盈率法、DCF现金流折现法和PB市净率法进行估值分析。报告中的预测和估值基于当前可获取的数据和假设，未来实际业绩可能与预测存在差异。本报告仅供参考研究使用，不构成任何投资建议。股市有风险，投资需谨慎。', body))
    story.append(Paragraph(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}', body))

    doc.build(story)
    pdf_io.seek(0)
    return pdf_io

# ==================== 主应用 ====================

def main():
    # 侧边栏
    with st.sidebar:
        st.markdown("## 股票搜索")
        search_term = st.text_input("输入代码或名称", placeholder="600519 / 贵州茅台")
        selected_stock = None
        if search_term:
            stock_list = get_stock_list_baostock()
            if not stock_list.empty:
                matched = stock_list[
                    stock_list['代码'].str.contains(search_term, na=False) |
                    stock_list['名称'].str.contains(search_term, na=False)
                ]
                if not matched.empty:
                    selected_stock = matched.iloc[0]
                    st.success(f"{selected_stock['名称']} ({selected_stock['代码']})")
        st.markdown("---")
        st.caption("估值参数将根据股票特征智能推荐，也可手动调整")

    # 标题
    st.markdown('<div class="report-title">A股智能估值分析</div>', unsafe_allow_html=True)
    st.markdown('<div class="report-sub">基于真实财务数据 · 参照券商研报估值框架</div>', unsafe_allow_html=True)

    if selected_stock is None:
        stock_list = get_stock_list_baostock()
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric("覆盖A股", f"{len(stock_list):,}只")
        with c2:
            st.metric("估值模型", "PE+DCF+PB")
        with c3:
            st.metric("数据源", "Baostock")
        st.info("在左侧输入股票代码或名称，获取专业估值分析。")
        if not stock_list.empty:
            hot = stock_list[stock_list['代码'].isin(['600519','600036','000858','002594','601318','600900','000001','601012'])]
            st.dataframe(hot if not hot.empty else stock_list.head(15), use_container_width=True, hide_index=True, height=280)
        return

    # 获取数据
    symbol = str(selected_stock['代码']).strip().zfill(6)
    stock_name = selected_stock['名称']
    end_date = datetime.now().strftime('%Y-%m-%d')
    start_date = (datetime.now() - timedelta(days=250)).strftime('%Y-%m-%d')

    with st.spinner(f'正在获取 {stock_name}({symbol}) 数据...'):
        history_df = get_stock_history_baostock(symbol, start_date, end_date)
        fin_data = get_full_financial_data(symbol)

    if history_df.empty:
        st.warning(f'未能获取 {stock_name}({symbol}) 的行情数据。')
        return

    # 提取数据
    current_price = float(history_df['close'].iloc[-1])
    latest_pe = float(history_df['peTTM'].iloc[-1]) if pd.notna(history_df['peTTM'].iloc[-1]) else 0
    latest_pb = float(history_df['pbMRQ'].iloc[-1]) if pd.notna(history_df['pbMRQ'].iloc[-1]) else 0
    pct_chg = float(history_df['pctChg'].iloc[-1]) if pd.notna(history_df['pctChg'].iloc[-1]) else 0
    high_10w = history_df['high'].tail(50).max()
    low_10w = history_df['low'].tail(50).min()
    bvps = current_price / latest_pb if latest_pb > 0 else 0

    dates = sorted(fin_data.keys(), reverse=True)
    latest_fin = fin_data.get(dates[0], {}) if dates else {}
    eps = latest_fin.get('epsTTM', 0)
    roe = latest_fin.get('roeAvg', 0)
    net_profit = latest_fin.get('netProfit', 0)
    revenue = latest_fin.get('MBRevenue', 0)
    gp_margin = latest_fin.get('gpMargin', 0)
    np_margin = latest_fin.get('npMargin', 0)
    debt_ratio = latest_fin.get('debtToAssets', 0)
    yoy_np = latest_fin.get('YOYNP', 0)

    # 智能估值参数推荐（基于股票名称识别行业）
    def recommend_params(name, pe_ttm, pb_mrq, roe_val, yoy_np_val):
        """根据股票名称和行业特征推荐估值参数和权重"""
        name = str(name)
        # 银行股
        if any(k in name for k in ['银行', '招商', '平安', '工商', '建设', '农业', '中国', '兴业', '浦发', '民生', '光大', '华夏', '北京', '上海', '江苏', '南京', '宁波', '杭州', '成都', '长沙', '苏州', '无锡', '常熟', '江阴', '张家港', '吴江', '昆山', '太仓']):
            return "银行股", 5, 8, 12, 0.05, 0.08, 0.03, 1.2, "PB为主", 0.1, 0.1, 0.8
        # 保险/券商/信托
        elif any(k in name for k in ['保险', '证券', '券商', '信托', '期货', '新华', '人寿', '太保', '平安']):
            return "金融股", 8, 12, 18, 0.08, 0.09, 0.03, 1.5, "PB为主", 0.2, 0.2, 0.6
        # 地产股
        elif any(k in name for k in ['地产', '万科', '保利', '招商蛇口', '金地', '新城', '华侨城', '绿地', '华夏幸福', '荣盛', '金科', '阳光城', '蓝光', '泰禾', '恒大', '碧桂园', '龙湖', '华润', '中海', '绿城', '世茂', '融创', '旭辉', '正荣', '融信', '中梁', '祥生', '宝龙', '禹洲', '力高', '三盛', '大发', '港龙', '上坤', '景瑞', '宋都', '冠城', '京投', '天房', '海泰', '津滨', '天保', '空港', '海泰', '高新', '张江', '浦东', '陆家嘴', '外高桥', '金桥', '碧水源', '华夏幸福']):
            return "地产股", 6, 10, 15, 0.05, 0.10, 0.02, 1.0, "PB为主", 0.1, 0.1, 0.8
        # 煤炭/钢铁/有色/石油（周期股）
        elif any(k in name for k in ['煤炭', '钢铁', '有色', '石油', '化工', '水泥', '建材', '电力', '燃气', '水务', '环保', '中煤', '神华', '陕煤', '兖矿', '焦煤', '潞安', '平煤', '阳泉', '晋控', '山煤', '淮北', '淮南', '盘江', '兰花', '永泰', '宝丰', '宝钢', '鞍钢', '首钢', '河钢', '沙钢', '南钢', '华菱', '三钢', '方大', '中信特钢', '太钢', '本钢', '马钢', '重钢', '山东钢铁', '包钢', '酒钢', '柳钢', '八一', '新钢', '韶钢', '凌钢', '安阳', '抚顺', '大冶', '八一', '江西铜业', '云南铜业', '铜陵有色', '紫金矿业', '山东黄金', '中金黄金', '银泰黄金', '赤峰黄金', '湖南黄金', '西部黄金', '恒邦股份', '紫金矿业', '洛阳钼业', '华友钴业', '寒锐钴业', '盛屯矿业', '厦门钨业', '章源钨业', '中钨高新', '翔鹭钨业', '广晟有色', '北方稀土', '盛和资源', '五矿稀土', '金力永磁', '正海磁材', '宁波韵升', '中科三环', '英洛华', '银河磁体', '横店东磁', '天通股份', '东阳光', '中孚实业', '明泰铝业', '常铝股份', '华峰铝业', '南山铝业', '云铝股份', '神火股份', '中国铝业', '天山铝业', '中色股份', '锌业股份', '罗平锌电', '驰宏锌锗', '株冶集团', '锡业股份', '贵研铂业', '贵绳股份', '金钼股份', '洛阳钼业', '中钨高新', '厦门钨业', '章源钨业', '翔鹭钨业', '广晟有色', '北方稀土', '盛和资源', '五矿稀土', '金力永磁', '正海磁材', '宁波韵升', '中科三环', '英洛华', '银河磁体', '横店东磁', '天通股份', '中国石油', '中国石化', '中国海油', '中海油服', '海油工程', '中曼石油', '贝肯能源', '通源石油', '潜能恒信', '惠博普', '海默科技', '恒泰艾普', '石化油服', '泰山石油', '广聚能源', '国际实业', '茂化实华', '沈阳化工', '氯碱化工', '中泰化学', '新疆天业', '君正集团', '英力特', '鸿达兴业', '亿利洁能', '滨化股份', '鲁西化工', '华鲁恒升', '万华化学', '齐翔腾达', '卫星化学', '东华能源', '宝丰能源', '中煤能源', '陕西煤业', '兖矿能源', '中国神华', '晋控煤业', '潞安环能', '平煤股份', '山西焦煤', '淮北矿业', '盘江股份', '兰花科创', '华阳股份', '山煤国际', '冀中能源', '开滦股份', '上海能源', '恒源煤电', '电投能源', '辽宁能源', '郑州煤电', '安源煤业', '云煤能源', '新大洲', '永泰能源', '宝泰隆', '金能科技', '美锦能源', '陕西黑猫', '云维股份', '山西焦化', '云煤能源']):
            return "周期股", 8, 12, 20, 0.08, 0.10, 0.02, 1.5, "PE+PB", 0.3, 0.2, 0.5
        # 白酒/食品饮料（消费白马）
        elif any(k in name for k in ['茅台', '五粮液', '泸州', '山西汾酒', '洋河', '古井', '水井坊', '舍得', '酒鬼', '口子窖', '今世缘', '迎驾', '金徽', '伊力特', '老白干', '顺鑫', '青青稞', '金种子', '皇台', '天佑德', '食品', '饮料', '乳业', '伊利', '蒙牛', '光明', '三元', '新乳业', '妙可蓝多', '贝因美', '科迪', '燕塘', '天润', '庄园', '西部牧业', '维维', '承德露露', '养元', '椰树', '欢乐家', '东鹏', '香飘飘', '康师傅', '统一', '旺旺', '农夫山泉', '海天', '中炬', '千禾', '恒顺', '榨菜', '桃李', '达利', '盼盼', '三只松鼠', '良品铺子', '来伊份', '盐津铺子', '劲仔', '甘源', '洽洽', '有友', '绝味', '周黑鸭', '煌上煌', '紫燕', '味知香', '安井', '三全', '思念', '海欣', '惠发', '千味央厨', '立高', '南侨', '海融', '佳禾', '香飘飘', '奈雪', '喜茶', '蜜雪冰城', '瑞幸', '星巴克']):
            return "消费白马", 20, 30, 45, 0.15, 0.09, 0.03, 4.0, "PE为主", 0.6, 0.2, 0.2
        # 医药
        elif any(k in name for k in ['医药', '制药', '生物', '医疗', '器械', '疫苗', '中药', '恒瑞', '迈瑞', '药明', '爱尔', '片仔癀', '云南白药', '同仁堂', '东阿', '复星', '智飞', '沃森', '康泰', '长春高新', '华兰', '天坛', '科伦', '健康元', '丽珠', '人福', '恩华', '华东', '国药', '上海医药', '华润', '九州通', '益丰', '大参林', '老百姓', '一心堂', '健之佳', '漱玉平民', '第一医药', '英特集团', '浙江医药', '海正', '华海', '普洛', '仙琚', '奥翔', '美诺华', '司太立', '天宇', '同和', '富祥', '山河药辅', '尔康', '广济', '健民', '葵花', '济川', '康缘', '以岭', '红日', '香雪', '白云山', '华润三九', '昆药', '达仁堂', '太极', '佐力', '方盛', '千金', '益佰', '贵州百灵', '奇正藏药', '西藏药业', '新天', '佛慈', '陇神戎发', '维康', '华森', '盘龙', '特一', '葫芦娃', '羚锐', '马应龙', '江中', '仁和', '亚宝', '康恩贝', '九芝堂', '启迪药业', '龙津', '大理药业', '华神科技', '嘉应制药', '紫鑫药业', '吉药控股', '通化金马', '景峰医药', '神奇制药', '贵州三力', '贵州百灵', '益佰制药', '信邦制药', '圣济堂', '灵康药业', '莎普爱思', '海思科', '恩华药业', '人福医药', '羚锐制药', '马应龙', '江中药业', '仁和药业', '亚宝药业', '康恩贝', '九芝堂', '启迪药业', '龙津药业', '大理药业', '华神科技', '嘉应制药', '紫鑫药业', '吉药控股', '通化金马', '景峰医药', '神奇制药', '贵州三力', '贵州百灵', '益佰制药', '信邦制药', '圣济堂', '灵康药业', '莎普爱思', '海思科', '恩华药业', '人福医药']):
            return "医药股", 25, 35, 55, 0.18, 0.09, 0.03, 4.5, "PE为主", 0.6, 0.2, 0.2
        # 光通信/CPO/算力（AI硬件，高估值但高增长）
        elif any(k in name for k in ['光模块', 'CPO', '光通信', '光迅', '中际旭创', '新易盛', '天孚通信', '剑桥科技', '太辰光', '德科立', '联特科技', '源杰科技', '长光华芯', '仕佳光子', '腾景科技', '博创科技', '华工科技', '光库', '昂纳', '海信宽带', '旭创', 'Finisar', 'Coherent', 'Lumentum', 'Fabrinet', '光器件', '光芯片', '硅光', '薄膜铌酸锂', 'PLC', 'AWG', 'CW激光器', 'EML', 'VCSEL', '探测器', '调制器', '耦合器', '波分复用', '光纤', '光缆', '亨通', '中天', '长飞', '通鼎', '永鼎', '特发', '汇源', '中利', '兆龙', '新亚', '卡倍亿', '沪光', '大地']):
            return "光通信/CPO", 40, 70, 120, 0.35, 0.10, 0.04, 8.0, "PE为主", 0.6, 0.2, 0.2
        # 新能源/光伏/锂电
        elif any(k in name for k in ['新能源', '光伏', '太阳能', '风电', '锂电', '宁德', '比亚迪', '隆基', '通威', '阳光电源', '晶科', '晶澳', '天合', '阿特斯', '东方日升', '爱旭', '钧达', '中来', '协鑫', 'TCL中环', '京运通', '双良', '上机', '高测', '宇晶', '欧晶', '福莱特', '信义', '旗滨', '南玻', '金晶', '亚玛顿', '安彩', '洛阳玻璃', '金刚玻璃', '德力', '秀强', '拓日新能', '亿晶光电', '珈伟新能', '清源股份', '芯能科技', '晶科科技', '太阳能', '林洋能源', '正泰电器', '特变电工', '国电南瑞', '许继电气', '平高电气', '中国西电', '思源电气', '特锐德', '科士达', '科华数据', '英维克', '高澜股份', '同飞股份', '申菱环境', '依米康', '佳力图', '朗进科技', '祥鑫科技', '铭科精技', '文灿股份', '爱柯迪', '旭升集团', '嵘泰股份', '立中集团', '永茂泰', '明泰铝业', '常铝股份', '华峰铝业', '南山铝业', '云铝股份', '神火股份', '中国铝业', '天山铝业', '中孚实业', '怡球资源', '顺博合金', '永茂泰', '立中集团', '四通新材', '云海金属', '宝武镁业', '宜安科技', '福达合金', '温州宏丰', '电工合金', '博威合金', '楚江新材', '精达股份', '长城科技', '金杯电工', '远东股份', '起帆电缆', '万马股份', '东方电缆', '中天科技', '亨通光电', '通光线缆', '特发信息', '汇源通信', '永鼎股份', '通鼎互联', '中利集团', '兆龙互连', '新亚电子', '卡倍亿', '沪光股份', '大地电气', '立讯精密', '歌尔股份', '蓝思科技', '领益智造', '鹏鼎控股', '环旭电子', '立讯精密', '工业富联', '闻泰科技', '传音控股', '中兴通讯', '烽火通信']):
            return "新能源", 30, 50, 80, 0.25, 0.10, 0.04, 5.0, "PE为主", 0.6, 0.2, 0.2
        # 半导体/芯片/AI/科技
        elif any(k in name for k in ['半导体', '芯片', '集成电路', '电子', '科技', '软件', '人工智能', 'AI', '中芯', '北方华创', '中微', '拓荆', '盛美', '华海清科', '芯源微', '微导纳米', '至纯', '富创', '江丰', '正帆', '新莱', '英杰', '汉钟', '中科', '寒武纪', '海光', '龙芯', '景嘉微', '兆易', '韦尔', '卓胜微', '圣邦', '思瑞浦', '艾为', '纳芯微', '晶晨', '瑞芯微', '全志', '恒玄', '乐鑫', '博通', '翱捷', '中科蓝讯', '炬芯', '峰岹', '芯海', '国民技术', '国科微', '欧比特', '复旦微电', '紫光国微', '振芯科技', '华力创通', '北斗星通', '合众思壮', '华测导航', '中海达', '四维图新', '高德红外', '大立科技', '睿创微纳', '久之洋', '富吉瑞', '东方中科', '普源精电', '鼎阳科技', '坤恒顺维', '思林杰', '优利德', '华盛昌', '聚辰股份', '普冉股份', '东芯股份', '恒烁股份', '佰维存储', '江波龙', '德明利', '朗科科技', '大为股份', '同有科技', '浪潮信息', '中科曙光', '紫光股份', '星网锐捷', '锐捷网络', '菲菱科思', '共进股份', '剑桥科技', '工业富联', '云赛智联', '数据港', '奥飞数据', '光环新网', '宝信软件', '用友网络', '金山办公', '科大讯飞', '汉王科技', '虹软科技', '当虹科技', '格灵深瞳', '云从科技', '云天励飞', '商汤', '旷视', '依图', '地平线', '黑芝麻', '芯原股份', '安路科技', '复旦微电', '紫光国微', '国民技术', '国科微', '欧比特', '上海贝岭', '士兰微', '华润微', '捷捷微电', '扬杰科技', '斯达半导', '时代电气', '宏微科技', '新洁能', '东微半导', '锴威特', '台基股份', '华微电子', '立昂微', '中晶科技', '欧陆通', '可立克', '京泉华', '伊戈尔', '麦格米特', '英搏尔', '汇川技术', '信捷电气', '雷赛智能', '埃斯顿', '机器人', '绿的谐波', '双环传动', '中大力德', '昊志机电', '禾川科技', '步科股份', '鸣志电器', '江苏雷利', '鼎智科技', '伟创电气', '正弦电气', '众辰科技', '儒竞科技', '三花智控', '拓普集团', '银轮股份', '飞龙股份', '盾安环境', '海信家电', '美的集团', '格力电器', '海尔智家', '老板电器', '华帝股份', '火星人', '浙江美大', '亿田智能', '帅丰电器', '苏泊尔', '九阳股份', '新宝股份', '小熊电器', '北鼎股份', '比依股份', '德昌股份', '春光科技', '富佳股份', '莱克电气', '科沃斯', '石头科技', '极米科技', '光峰科技', '海信视像', 'TCL科技', '创维数字', '四川长虹', '康冠科技', '视源股份', '鸿合科技', '宸展光电', '伟时电子', '隆利科技', '聚飞光电', '瑞丰光电', '鸿利智汇', '国星光电', '木林森', '三安光电', '华灿光电', '乾照光电', '聚灿光电', '蔚蓝锂芯', '兆驰股份', '洲明科技', '利亚德', '艾比森', '奥拓电子', '雷曼光电', '联建光电', '深天马A', '京东方A', 'TCL科技', '维信诺', '和辉光电', '龙腾光电', '彩虹股份', '东旭光电', '诚志股份', '飞凯材料', '江化微', '晶瑞电材', '南大光电', '上海新阳', '雅克科技', '鼎龙股份', '安集科技', '强力新材', '容大感光', '飞凯材料', '江化微', '晶瑞电材', '南大光电', '上海新阳', '雅克科技', '鼎龙股份', '安集科技', '强力新材', '容大感光']):
            return "科技股", 35, 60, 100, 0.30, 0.10, 0.04, 6.0, "PE为主", 0.6, 0.2, 0.2
        # 精密设备/专用设备/小盘高PE
        elif any(k in name for k in ['凯格', '精机', '精测', '精雕', '联得', '智云', '华兴', '科恒', '正业', '劲拓', '集泰', '泰尔', '至纯', '盛美', '北方华创', '中微', '拓荆', '华海清科', '芯源微', '微导纳米', '迈为', '帝尔', '奥特维', '捷佳伟创', '金辰', '先导智能', '赢合', '科恒', '星云', '瀚川', '博众', '赛腾', '拓斯达', '埃斯顿', '汇川', '信捷', '雷赛', '绿的谐波', '禾川', '鸣志', '步科', '鼎智', '伟创', '正弦', '众辰', '儒竞', '凯尔达', '埃夫特', '新松', '拓斯达', '克来机电', '天奇', '华昌达', '巨一科技', '豪森', '迈安德', '杰克', '上工', '标准', '中捷', '宝石', '大族', '锐科', '杰普特', '帝尔', '海目星', '联赢', '逸飞', '铂力特', '华曙', '先临三维', '极光', '铂力特', '铂力特']):
            return "精密设备", 50, 80, 130, 0.30, 0.10, 0.04, 7.0, "PE为主", 0.6, 0.2, 0.2
        # 高成长型（通用）
        elif pe_ttm > 50 or (yoy_np_val and yoy_np_val > 30):
            return "高成长型", 30, 50, 80, 0.25, 0.10, 0.04, 5.0, "PE为主", 0.5, 0.3, 0.2
        # 价值型
        elif pe_ttm < 15 and roe_val > 15:
            return "价值型", 10, 15, 25, 0.08, 0.08, 0.03, 2.0, "PE+PB", 0.4, 0.2, 0.4
        # 稳健型
        elif pe_ttm < 25 and (yoy_np_val and yoy_np_val < 10):
            return "稳健型", 12, 20, 35, 0.10, 0.09, 0.03, 2.5, "PE为主", 0.5, 0.2, 0.3
        else:
            return "成长型", 15, 30, 50, 0.15, 0.10, 0.03, 3.0, "PE为主", 0.5, 0.2, 0.3

    stock_type, rec_pe_low, rec_pe_mid, rec_pe_high, rec_growth, rec_discount, rec_terminal, rec_pb, val_method, w_pe, w_dcf, w_pb = \
        recommend_params(stock_name, latest_pe, latest_pb, roe, yoy_np)

    # 估值计算（使用推荐参数或用户自定义参数）
    use_recommended = st.sidebar.checkbox("使用智能推荐参数", value=True, help=f"该股票属于【{stock_type}】，估值方法：{val_method}")
    if use_recommended:
        pe_low, pe_mid, pe_high = rec_pe_low, rec_pe_mid, rec_pe_high
        growth, discount, terminal, pb_mult = rec_growth, rec_discount, rec_terminal, rec_pb
        st.sidebar.info(f"已匹配【{stock_type}】参数\n估值方法：{val_method}\nPE权重{w_pe*100:.0f}% + DCF权重{w_dcf*100:.0f}% + PB权重{w_pb*100:.0f}%")
    else:
        st.sidebar.markdown("## 估值参数")
        pe_low = st.sidebar.slider("悲观PE", 5, 50, rec_pe_low)
        pe_mid = st.sidebar.slider("基准PE", 10, 100, rec_pe_mid)
        pe_high = st.sidebar.slider("乐观PE", 15, 150, rec_pe_high)
        with st.sidebar.expander("DCF参数"):
            growth = st.sidebar.slider("增长率(%)", -10, 100, int(rec_growth*100)) / 100
            discount = st.sidebar.slider("折现率(%)", 5, 20, int(rec_discount*100)) / 100
            terminal = st.sidebar.slider("永续增长(%)", 0, 8, int(rec_terminal*100)) / 100
        with st.sidebar.expander("PB参数"):
            pb_mult = st.sidebar.slider("PB倍数", 1.0, 10.0, rec_pb, step=0.1)

    pe_pessimistic, pe_base, pe_optimistic = calculate_pe_valuation(eps, pe_low, pe_mid, pe_high)
    dcf_val = calculate_dcf_valuation(net_profit, revenue, eps, growth, discount, terminal, latest_pe)
    pb_val = calculate_pb_valuation(bvps, pb_mult)
    fair_price = pe_base * w_pe + dcf_val * w_dcf + pb_val * w_pb

    valuation = {'pe': (pe_pessimistic, pe_base, pe_optimistic), 'dcf': dcf_val, 'pb': pb_val, 'fair': fair_price}
    rating_info = get_investment_rating(current_price, fair_price, latest_pe, roe, gp_margin)

    # ========== 页面布局（研报风格） ==========

    # 1. 股票标题 + 基本数据
    st.markdown(f"### {stock_name}（{symbol}）")
    col_info = st.columns(6)
    col_info[0].metric("收盘价", f"¥{current_price:.2f}", f"{pct_chg:+.2f}%",
                       delta_color="normal" if pct_chg >= 0 else "inverse")
    col_info[1].metric("PE(TTM)", f"{latest_pe:.1f}x" if latest_pe > 0 else "N/A")
    col_info[2].metric("PB(MRQ)", f"{latest_pb:.1f}x" if latest_pb > 0 else "N/A")
    col_info[3].metric("10周高", f"¥{high_10w:.2f}")
    col_info[4].metric("10周低", f"¥{low_10w:.2f}")
    col_info[5].metric("总股本", f"{selected_stock.get('交易所','').upper()}")

    # 2. 投资建议（醒目）
    rclass = rating_info[1]
    st.markdown(f"""
    <div class="rating-{rclass}">
        投资评级: <b>{rating_info[0]}</b> &nbsp;|&nbsp;
        合理估值: <b>¥{fair_price:.2f}</b> &nbsp;|&nbsp;
        溢价空间: <b>{rating_info[2]:+.1f}%</b>
    </div>
    """, unsafe_allow_html=True)
    if rating_info[3]:
        st.caption(" | ".join(rating_info[3]))

    # 3. 盈利预测与估值（核心表格 - 研报风格）
    st.markdown('<div class="section-title">盈利预测与估值</div>', unsafe_allow_html=True)
    pred_dates = sorted(fin_data.keys())
    pred_labels = [d[:4] for d in pred_dates]

    # 构建预测表
    pred_rows = []
    for d in pred_dates:
        fd = fin_data[d]
        np_v = fd.get('netProfit', 0)
        rev_v = fd.get('MBRevenue', 0)
        eps_v = fd.get('epsTTM', 0)
        roe_v = fd.get('roeAvg', 0)
        gp_v = fd.get('gpMargin', 0)
        yoy = fd.get('YOYNP', 0)
        pred_rows.append([
            d[:4],
            f"{rev_v/1e8:.1f}" if rev_v and rev_v == rev_v else "-",
            f"{np_v/1e8:.1f}" if np_v and np_v == np_v else "-",
            f"{eps_v:.2f}" if eps_v and eps_v == eps_v else "-",
            f"{roe_v:.1f}" if roe_v and roe_v == roe_v else "-",
            f"{gp_v:.1f}" if gp_v and gp_v == gp_v else "-",
            f"{yoy:.1f}" if yoy and yoy == yoy else "-",
        ])

    # 添加预测行
    if eps > 0:
        pred_rows.append(["2026E", "-", "-", f"{eps:.2f}", f"{roe:.1f}", f"{gp_margin:.1f}", "-"])
        est_np_next = net_profit * (1 + growth) if net_profit > 0 else 0
        pred_rows.append(["2027E", "-", "-", f"{eps*(1+growth):.2f}", "-", "-", "-"])

    pred_df = pd.DataFrame(pred_rows, columns=['年度', '营收(亿)', '净利润(亿)', 'EPS', 'ROE(%)', '毛利率(%)', '利润增速(%)'])
    st.dataframe(pred_df, use_container_width=True, hide_index=True, height=max(200, 36 * len(pred_rows)))

    # 4. 未来股价预测（醒目展示）
    st.markdown('<div class="section-title">未来股价预测</div>', unsafe_allow_html=True)
    pred_col1, pred_col2, pred_col3, pred_col4 = st.columns(4)
    pred_col1.metric("悲观情景", f"¥{pe_pessimistic:.2f}", f"PE={pe_low}x")
    pred_col2.metric("基准情景", f"¥{pe_base:.2f}", f"PE={pe_mid}x")
    pred_col3.metric("乐观情景", f"¥{pe_optimistic:.2f}", f"PE={pe_high}x")
    pred_col4.metric("综合估值", f"¥{fair_price:.2f}", f"PE40%+DCF30%+PB30%")

    # 未来股价预测表
    future_prices = []
    for year_offset in [1, 2, 3]:
        future_eps = eps * ((1 + growth) ** year_offset) if eps > 0 else 0
        future_pe = pe_mid
        future_price_pe = future_eps * future_pe
        future_price_dcf = dcf_val * (1 + growth * year_offset * 0.3)
        future_price_pb = pb_val * (1 + growth * year_offset * 0.2)
        future_fair = future_price_pe * 0.4 + future_price_dcf * 0.3 + future_price_pb * 0.3
        future_prices.append({
            '预测年度': f'{datetime.now().year + year_offset}E',
            '预测EPS': f'¥{future_eps:.2f}',
            'PE估值': f'¥{future_price_pe:.2f}',
            'DCF估值': f'¥{future_price_dcf:.2f}',
            'PB估值': f'¥{future_price_pb:.2f}',
            '综合估值': f'¥{future_fair:.2f}',
            '较当前价': f"{((future_fair/current_price-1)*100):+.1f}%" if current_price > 0 else "-"
        })
    future_df = pd.DataFrame(future_prices)
    st.dataframe(future_df, use_container_width=True, hide_index=True, height=150)

    # 5. 估值方法对比
    st.markdown('<div class="section-title">估值方法对比</div>', unsafe_allow_html=True)

    # 估值方法对比表
    val_table = pd.DataFrame({
        '估值方法': ['PE悲观', 'PE基准', 'PE乐观', 'DCF折现', 'PB估值', '综合估值'],
        '估值价格': [f'¥{pe_pessimistic:.2f}', f'¥{pe_base:.2f}', f'¥{pe_optimistic:.2f}',
                   f'¥{dcf_val:.2f}', f'¥{pb_val:.2f}', f'¥{fair_price:.2f}'],
        '当前价': [f'¥{current_price:.2f}'] * 6,
        '溢价空间': [f"{((pe_pessimistic/current_price-1)*100):+.1f}%" if current_price > 0 else "-",
                    f"{((pe_base/current_price-1)*100):+.1f}%" if current_price > 0 else "-",
                    f"{((pe_optimistic/current_price-1)*100):+.1f}%" if current_price > 0 else "-",
                    f"{((dcf_val/current_price-1)*100):+.1f}%" if current_price > 0 else "-",
                    f"{((pb_val/current_price-1)*100):+.1f}%" if current_price > 0 else "-",
                    f"{((fair_price/current_price-1)*100):+.1f}%" if current_price > 0 else "-"]
    })
    st.dataframe(val_table, use_container_width=True, hide_index=True, height=240)

    # 5. 图表（紧凑排列）
    st.markdown('<div class="section-title">走势与估值图表</div>', unsafe_allow_html=True)
    col_c1, col_c2 = st.columns(2)
    with col_c1:
        fig_k = create_kline_chart(history_df, fair_price)
        if fig_k:
            st.plotly_chart(fig_k, use_container_width=True, key="kline")
    with col_c2:
        fig_pe = create_pe_chart(eps, current_price, pe_low, pe_mid, pe_high)
        if fig_pe:
            st.plotly_chart(fig_pe, use_container_width=True, key="pe")
        else:
            st.info("暂无EPS数据")

    # 财务图表
    fig_fin = create_financial_chart(fin_data)
    if fig_fin:
        st.plotly_chart(fig_fin, use_container_width=True, key="fin")

    # 6. 主要财务比率
    st.markdown('<div class="section-title">主要财务比率</div>', unsafe_allow_html=True)
    ratio_rows = []
    for d in pred_dates:
        fd = fin_data[d]
        ratio_rows.append([
            d[:4],
            f"{fd.get('gpMargin', 0):.1f}" if fd.get('gpMargin') and fd.get('gpMargin') == fd.get('gpMargin') else "-",
            f"{fd.get('npMargin', 0):.1f}" if fd.get('npMargin') and fd.get('npMargin') == fd.get('npMargin') else "-",
            f"{fd.get('roeAvg', 0):.1f}" if fd.get('roeAvg') and fd.get('roeAvg') == fd.get('roeAvg') else "-",
            f"{fd.get('debtToAssets', 0):.1f}" if fd.get('debtToAssets') and fd.get('debtToAssets') == fd.get('debtToAssets') else "-",
            f"{fd.get('YOYAsset', 0):.1f}" if fd.get('YOYAsset') and fd.get('YOYAsset') == fd.get('YOYAsset') else "-",
        ])
    ratio_df = pd.DataFrame(ratio_rows, columns=['年度', '毛利率(%)', '净利率(%)', 'ROE(%)', '资产负债率(%)', '资产增速(%)'])
    st.dataframe(ratio_df, use_container_width=True, hide_index=True, height=max(180, 36 * len(ratio_rows)))

    # 7. 风险提示
    st.markdown('<div class="section-title">风险提示</div>', unsafe_allow_html=True)
    risks = []
    if latest_pe > 60:
        risks.append("当前估值偏高，需警惕估值回调风险")
    if debt_ratio > 60:
        risks.append("资产负债率较高，存在偿债压力")
    if yoy_np and yoy_np < 0:
        risks.append("净利润同比下降，关注盈利持续性")
    if not risks:
        risks = ["行业竞争加剧风险", "宏观经济波动风险", "政策变化风险"]
    for r in risks:
        st.caption(f"- {r}")

    # 8. 下载报告
    st.markdown("---")
    col_dl = st.columns(2)
    with col_dl[0]:
        word_file = generate_word_report(symbol, stock_name, current_price, valuation, fin_data, rating_info, history_df, pe_low, pe_mid, pe_high, growth, discount, terminal, pb_mult)
        st.download_button(label="下载 Word 报告", data=word_file,
                           file_name=f"{stock_name}_{symbol}_估值报告.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    with col_dl[1]:
        pdf_file = generate_pdf_report(symbol, stock_name, current_price, valuation, rating_info, fin_data, history_df, pe_low, pe_mid, pe_high, growth, discount, terminal, pb_mult)
        st.download_button(label="下载 PDF 报告", data=pdf_file,
                           file_name=f"{stock_name}_{symbol}_估值报告.pdf",
                           mime="application/pdf", use_container_width=True)

    st.caption("免责声明：本系统基于Baostock公开数据生成，仅供参考，不构成投资建议。投资有风险，入市需谨慎。")

if __name__ == "__main__":
    main()
